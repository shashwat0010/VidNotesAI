import io
import os
import re
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from app.models.models import NoteOutput

class ExportService:
    @staticmethod
    def generate_markdown(note: NoteOutput, title: str, keyframes: list = None) -> str:
        from app.services.llm import llm_service
        sanitized = llm_service._deterministic_sanitize_notes_dict({
            "summary_exec": note.summary_exec or "",
            "summary_detailed": note.summary_detailed or "",
            "revision_notes": note.revision_notes or "",
            "takeaways": note.takeaways or "",
            "glossary": note.glossary or ""
        })

        md = []
        md.append(f"# {title} - Study Notes")
        md.append("\n## Executive Summary\n")
        md.append(sanitized["summary_exec"])
        
        md.append("\n## Detailed Lecture Notes\n")
        md.append(sanitized["summary_detailed"])
        
        md.append("\n## Revision & Review Checklist\n")
        md.append(sanitized["revision_notes"])
        
        md.append("\n## Key Takeaways\n")
        md.append(sanitized["takeaways"])
        
        md.append("\n## Glossary of Terms\n")
        md.append(sanitized["glossary"])
        
        # Check if keyframes need to be appended
        if keyframes:
            detailed_txt = sanitized["summary_detailed"]
            missing_kfs = [
                kf for kf in keyframes
                if (kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url'))
                and (kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url')) not in detailed_txt
            ]
            if missing_kfs:
                md.append("\n## Extracted Keyframe Slides\n")
                for kf in missing_kfs:
                    url = kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url')
                    ts = kf.timestamp if hasattr(kf, 'timestamp') else kf.get('timestamp', 0)
                    mins = int(ts // 60)
                    secs = int(ts % 60)
                    md.append(f"### Slide at {mins:02d}:{secs:02d}")
                    md.append(f"![Slide at {mins:02d}:{secs:02d}]({url})\n")

        if note.flashcards:
            md.append("\n## Flashcards\n")
            for idx, card in enumerate(note.flashcards, 1):
                md.append(f"**Q{idx}:** {card.get('question')}")
                md.append(f"**A{idx}:** {card.get('answer')}\n")
            
        if note.mcqs:
            md.append("\n## Multiple Choice Quiz\n")
            for idx, mcq in enumerate(note.mcqs, 1):
                md.append(f"**Question {idx}:** {mcq.get('question')}")
                for opt in mcq.get('options', []):
                    md.append(f"- {opt}")
                md.append(f"\n*Correct Answer: {mcq.get('answer')}*")
                md.append(f"*Explanation: {mcq.get('explanation')}*\n")

        if note.mindmap:
            md.append("\n## Mermaid Mind Map\n")
            md.append("```mermaid")
            md.append(note.mindmap)
            md.append("```")

        return "\n".join(md)

    @staticmethod
    def generate_docx(note: NoteOutput, title: str, keyframes: list = None) -> bytes:
        doc = Document()
        
        # Configure margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Base style configuration
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Title
        t = doc.add_heading(level=0)
        run = t.add_run(f"{title} - Study Notes")
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(31, 41, 55)

        uploads_base = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "uploads")

        def _get_clean_s3_key(url: str) -> str:
            if not url:
                return ""
            clean = url.split("?")[0]
            if ".amazonaws.com/" in clean:
                clean = clean.split(".amazonaws.com/")[1]
            clean = clean.replace("/uploads/", "").replace("/vidnotes-storage/", "").replace("uploads/", "").replace("vidnotes-storage/", "").lstrip("/")
            return clean

        def _get_docx_img_stream(url: str):
            if not url:
                return None
            key = _get_clean_s3_key(url)
            
            # 1. Try local filesystem cache first
            candidates = [
                os.path.join(uploads_base, key),
                os.path.join(uploads_base, "keyframes", key),
            ]
            for fp in candidates:
                if os.path.exists(fp) and os.path.isfile(fp):
                    try:
                        with open(fp, "rb") as f:
                            return io.BytesIO(f.read())
                    except Exception:
                        pass
            
            # 2. Try authenticated AWS S3 SDK
            if s3_service.s3 and settings.S3_BUCKET_NAME and key:
                try:
                    obj = s3_service.s3.get_object(Bucket=settings.S3_BUCKET_NAME, Key=key)
                    return io.BytesIO(obj['Body'].read())
                except Exception as s3_err:
                    print(f"[DOCX S3 Fetch Notice]: {s3_err}")

            # 3. Try HTTP request
            if url.startswith("http://") or url.startswith("https://"):
                try:
                    resp = requests.get(url, timeout=5)
                    if resp.status_code == 200:
                        return io.BytesIO(resp.content)
                except Exception:
                    pass
            return None

        def _parse_md_docx(text: str):
            if not isinstance(text, str):
                return
            for line in text.splitlines():
                l_str = line.strip()
                if not l_str:
                    continue

                # Check images
                img_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
                matches = list(re.finditer(img_pattern, l_str))
                if matches:
                    last_idx = 0
                    for match in matches:
                        start_pos, end_pos = match.span()
                        txt_before = l_str[last_idx:start_pos].strip()
                        if txt_before:
                            doc.add_paragraph(txt_before)
                        caption_text = match.group(1)
                        img_url = match.group(2)
                        stream = _get_docx_img_stream(img_url)
                        if stream:
                            try:
                                doc.add_picture(stream, width=Inches(4.5))
                                if caption_text:
                                    cap_p = doc.add_paragraph()
                                    cap_run = cap_p.add_run(caption_text)
                                    cap_run.italic = True
                                    cap_run.font.size = Pt(9)
                            except Exception as e:
                                print(f"[DOCX] Add picture failed: {e}")
                        last_idx = end_pos
                    txt_after = l_str[last_idx:].strip()
                    if txt_after:
                        doc.add_paragraph(txt_after)
                    continue

                # Headings
                if l_str.startswith("#### "):
                    p = doc.add_heading(l_str[5:], level=4)
                    p.style.font.color.rgb = RGBColor(75, 85, 99)
                elif l_str.startswith("### "):
                    p = doc.add_heading(l_str[4:], level=3)
                    p.style.font.color.rgb = RGBColor(55, 65, 81)
                elif l_str.startswith("## "):
                    p = doc.add_heading(l_str[3:], level=2)
                    p.style.font.color.rgb = RGBColor(31, 41, 55)
                elif l_str.startswith("# "):
                    p = doc.add_heading(l_str[2:], level=1)
                    p.style.font.color.rgb = RGBColor(30, 64, 175)
                elif l_str.startswith("- ") or l_str.startswith("* "):
                    doc.add_paragraph(l_str[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s+(.+)', l_str):
                    match = re.match(r'^\d+\.\s+(.+)', l_str)
                    doc.add_paragraph(match.group(1), style='List Number')
                else:
                    doc.add_paragraph(l_str)

        from app.services.llm import llm_service
        sanitized = llm_service._deterministic_sanitize_notes_dict({
            "summary_exec": note.summary_exec or "",
            "summary_detailed": note.summary_detailed or "",
            "revision_notes": note.revision_notes or "",
            "takeaways": note.takeaways or "",
            "glossary": note.glossary or ""
        })

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        _parse_md_docx(sanitized["summary_exec"])

        # Detailed summary
        doc.add_heading("Detailed Lecture Notes", level=1)
        _parse_md_docx(sanitized["summary_detailed"])

        # Revision Guide
        doc.add_heading("Revision Guide", level=1)
        _parse_md_docx(sanitized["revision_notes"])

        # Takeaways
        doc.add_heading("Key Takeaways", level=1)
        _parse_md_docx(sanitized["takeaways"])

        # Glossary
        doc.add_heading("Glossary of Terms", level=1)
        _parse_md_docx(sanitized["glossary"])

        # Append keyframes if not already inside detailed notes
        if keyframes:
            detailed_txt = sanitized["summary_detailed"]
            missing_kfs = [
                kf for kf in keyframes
                if (kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url'))
                and (kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url')) not in detailed_txt
            ]
            if missing_kfs:
                doc.add_heading("Extracted Keyframe Slides", level=1)
                for kf in missing_kfs:
                    url = kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url')
                    ts = kf.timestamp if hasattr(kf, 'timestamp') else kf.get('timestamp', 0)
                    mins = int(ts // 60)
                    secs = int(ts % 60)
                    doc.add_heading(f"Slide at {mins:02d}:{secs:02d}", level=2)
                    stream = _get_docx_img_stream(url)
                    if stream:
                        try:
                            doc.add_picture(stream, width=Inches(4.5))
                        except Exception as e:
                            print(f"[DOCX] Slide picture add notice: {e}")

        # Flashcards
        if note.flashcards:
            doc.add_heading("Flashcards", level=1)
            for idx, card in enumerate(note.flashcards, 1):
                p = doc.add_paragraph()
                p.add_run(f"Q{idx}: ").bold = True
                p.add_run(f"{card.get('question')}\n")
                p.add_run(f"A{idx}: ").bold = True
                p.add_run(f"{card.get('answer')}")

        # Quiz
        if note.mcqs:
            doc.add_heading("Multiple Choice Quiz", level=1)
            for idx, mcq in enumerate(note.mcqs, 1):
                p = doc.add_paragraph()
                p.add_run(f"Question {idx}: ").bold = True
                p.add_run(f"{mcq.get('question')}\n")
                for o in mcq.get('options', []):
                    p.add_run(f"  [ ] {o}\n")
                p.add_run("Correct Answer: ").bold = True
                p.add_run(f"{mcq.get('answer')}\n")
                p.add_run("Explanation: ").italic = True
                p.add_run(f"{mcq.get('explanation')}")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    @staticmethod
    def generate_pdf(note: NoteOutput, title: str, keyframes: list = None) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=60,
            leftMargin=60,
            topMargin=60,
            bottomMargin=60
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'DocTitle', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=22, leading=28,
            textColor=colors.HexColor('#111827'), spaceAfter=20
        )
        h1_style = ParagraphStyle(
            'H1', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=16, leading=20,
            textColor=colors.HexColor('#1e40af'),
            spaceBefore=18, spaceAfter=8, keepWithNext=True
        )
        h2_style = ParagraphStyle(
            'H2', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=13, leading=17,
            textColor=colors.HexColor('#1f2937'),
            spaceBefore=12, spaceAfter=6, keepWithNext=True
        )
        h3_style = ParagraphStyle(
            'H3', parent=styles['Normal'],
            fontName='Helvetica-BoldOblique', fontSize=11, leading=15,
            textColor=colors.HexColor('#374151'),
            spaceBefore=8, spaceAfter=4, keepWithNext=True
        )
        body_style = ParagraphStyle(
            'Body', parent=styles['Normal'],
            fontName='Helvetica', fontSize=10, leading=15,
            textColor=colors.HexColor('#374151'), spaceAfter=4
        )
        bullet_style = ParagraphStyle(
            'Bullet', parent=body_style,
            leftIndent=18, bulletIndent=6, spaceAfter=3
        )
        caption_style = ParagraphStyle(
            'Caption', parent=styles['Normal'],
            fontName='Helvetica-Oblique', fontSize=8, leading=10,
            textColor=colors.HexColor('#6b7280'),
            alignment=1, spaceAfter=8  # centered
        )
        code_style = ParagraphStyle(
            'Code', parent=styles['Normal'],
            fontName='Courier', fontSize=8.5, leading=12,
            textColor=colors.HexColor('#e2e8f0'),
            spaceBefore=2, spaceAfter=2
        )

        uploads_base = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "uploads")

        def _inline_md(text: str) -> str:
            if not text:
                return ""
            text = str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            text = re.sub(r'`(.+?)`', r'<font name="Courier">\1</font>', text)
            return text

        def _fetch_image(url: str):
            if not url:
                return None
            
            clean = url.split("?")[0]
            if ".amazonaws.com/" in clean:
                clean = clean.split(".amazonaws.com/")[1]
            key = clean.replace("/uploads/", "").replace("/vidnotes-storage/", "").replace("uploads/", "").replace("vidnotes-storage/", "").lstrip("/")

            # 1. Check local cached file first
            candidates = [
                os.path.join(uploads_base, key),
                os.path.join(uploads_base, "keyframes", key),
            ]
            for fp in candidates:
                if os.path.exists(fp) and os.path.isfile(fp):
                    try:
                        with open(fp, "rb") as f:
                            img_buf = io.BytesIO(f.read())
                        return RLImage(img_buf, width=4.5 * inch, height=2.8 * inch, kind='proportional')
                    except Exception as e:
                        print(f"[PDF] Local image read notice: {e}")

            # 2. Fetch directly from authenticated AWS S3
            if s3_service.s3 and settings.S3_BUCKET_NAME and key:
                try:
                    obj = s3_service.s3.get_object(Bucket=settings.S3_BUCKET_NAME, Key=key)
                    img_buf = io.BytesIO(obj['Body'].read())
                    return RLImage(img_buf, width=4.5 * inch, height=2.8 * inch, kind='proportional')
                except Exception as s3_err:
                    print(f"[PDF S3 Fetch Notice]: {s3_err}")

            # 3. HTTP fetch fallback
            if url.startswith("http://") or url.startswith("https://"):
                try:
                    resp = requests.get(url, timeout=5)
                    if resp.status_code == 200:
                        img_buf = io.BytesIO(resp.content)
                        return RLImage(img_buf, width=4.5 * inch, height=2.8 * inch, kind='proportional')
                except Exception as e:
                    print(f"[PDF] Could not fetch image {url}: {e}")
            return None

        def _parse_markdown(text: str, story: list):
            lines = str(text or "").splitlines()
            i = 0
            while i < len(lines):
                line = lines[i]

                # Code block: ```lang ... ```
                if line.strip().startswith("```"):
                    code_lines = []
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith("```"):
                        code_lines.append(lines[i])
                        i += 1
                    if i < len(lines):
                        i += 1
                    code_text = "\n".join(code_lines)
                    code_html = str(code_text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
                    code_p = Paragraph(f"<font name='Courier' size=8.5 color='#e2e8f0'>{code_html}</font>", code_style)
                    code_table = Table([[code_p]], colWidths=[letter[0] - 120])
                    code_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#0f172a')),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ('LEFTPADDING', (0, 0), (-1, -1), 10),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#1e293b')),
                    ]))
                    story.append(Spacer(1, 4))
                    story.append(code_table)
                    story.append(Spacer(1, 6))
                    continue

                # Image: ![caption](url)
                img_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if img_match:
                    caption_text = img_match.group(1)
                    img_url = img_match.group(2)
                    img = _fetch_image(img_url)
                    if img:
                        story.append(Spacer(1, 6))
                        story.append(img)
                        if caption_text:
                            story.append(Paragraph(caption_text, caption_style))
                        story.append(Spacer(1, 6))
                    i += 1
                    continue

                # Headings
                if re.match(r'^####\s+(.+)', line):
                    text_content = re.match(r'^####\s+(.+)', line).group(1)
                    story.append(Paragraph(_inline_md(text_content), h3_style))
                    i += 1
                    continue

                if re.match(r'^###\s+(.+)', line):
                    text_content = re.match(r'^###\s+(.+)', line).group(1)
                    story.append(Paragraph(_inline_md(text_content), h3_style))
                    i += 1
                    continue

                if re.match(r'^##\s+(.+)', line):
                    text_content = re.match(r'^##\s+(.+)', line).group(1)
                    story.append(Paragraph(_inline_md(text_content), h2_style))
                    i += 1
                    continue

                if re.match(r'^#\s+(.+)', line):
                    text_content = re.match(r'^#\s+(.+)', line).group(1)
                    story.append(Paragraph(_inline_md(text_content), h1_style))
                    i += 1
                    continue

                # Horizontal rule
                if re.match(r'^---+$', line.strip()):
                    story.append(Spacer(1, 6))
                    t = Table([['']], colWidths=[letter[0] - 120])
                    t.setStyle(TableStyle([
                        ('LINEABOVE', (0, 0), (-1, -1), 0.5, colors.HexColor('#e5e7eb')),
                        ('TOPPADDING', (0, 0), (-1, -1), 0),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 6))
                    i += 1
                    continue

                # Bullet items
                bullet_match = re.match(r'^[-*]\s+(.+)', line)
                if bullet_match:
                    text_content = bullet_match.group(1)
                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{_inline_md(text_content)}", bullet_style))
                    i += 1
                    continue

                # Numbered list
                num_match = re.match(r'^(\d+)\.\s+(.+)', line)
                if num_match:
                    num = num_match.group(1)
                    text_content = num_match.group(2)
                    story.append(Paragraph(f"{num}. {_inline_md(text_content)}", bullet_style))
                    i += 1
                    continue

                # Empty line
                if not line.strip():
                    story.append(Spacer(1, 5))
                    i += 1
                    continue

                # Normal paragraph
                story.append(Paragraph(_inline_md(line), body_style))
                i += 1

        story = []

        # Title
        story.append(Paragraph(f"{title}", title_style))
        story.append(Paragraph("Study Notes", ParagraphStyle('Sub', parent=body_style,
            fontName='Helvetica-Oblique', fontSize=11, textColor=colors.HexColor('#6b7280'), spaceAfter=20)))
        story.append(Spacer(1, 12))

        from app.services.llm import llm_service
        sanitized = llm_service._deterministic_sanitize_notes_dict({
            "summary_exec": note.summary_exec or "",
            "summary_detailed": note.summary_detailed or "",
            "revision_notes": note.revision_notes or "",
            "takeaways": note.takeaways or "",
            "glossary": note.glossary or ""
        })

        # Sections
        story.append(Paragraph("Executive Summary", h1_style))
        _parse_markdown(sanitized["summary_exec"], story)
        story.append(Spacer(1, 10))

        story.append(Paragraph("Detailed Lecture Notes", h1_style))
        _parse_markdown(sanitized["summary_detailed"], story)
        story.append(Spacer(1, 10))

        story.append(Paragraph("Revision Guide", h1_style))
        _parse_markdown(sanitized["revision_notes"], story)
        story.append(Spacer(1, 10))

        story.append(Paragraph("Key Takeaways", h1_style))
        _parse_markdown(sanitized["takeaways"], story)
        story.append(Spacer(1, 10))

        story.append(Paragraph("Glossary of Terms", h1_style))
        _parse_markdown(sanitized["glossary"], story)
        story.append(Spacer(1, 10))

        # Keyframes slides
        if keyframes:
            detailed_txt = sanitized["summary_detailed"]
            missing_kfs = [
                kf for kf in keyframes
                if (kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url'))
                and (kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url')) not in detailed_txt
            ]
            if missing_kfs:
                story.append(Paragraph("Extracted Keyframe Slides", h1_style))
                for kf in missing_kfs:
                    url = kf.s3_url if hasattr(kf, 's3_url') else kf.get('s3_url')
                    ts = kf.timestamp if hasattr(kf, 'timestamp') else kf.get('timestamp', 0)
                    mins = int(ts // 60)
                    secs = int(ts % 60)
                    story.append(Paragraph(f"Slide at {mins:02d}:{secs:02d}", h2_style))
                    img = _fetch_image(url)
                    if img:
                        story.append(Spacer(1, 4))
                        story.append(img)
                        story.append(Spacer(1, 4))

        # Flashcards
        if note.flashcards:
            story.append(Paragraph("Flashcards", h1_style))
            for idx, card in enumerate(note.flashcards, 1):
                q = _inline_md(str(card.get('question', '')))
                a = _inline_md(str(card.get('answer', '')))
                story.append(Paragraph(f"<b>Q{idx}:</b> {q}", body_style))
                story.append(Paragraph(f"<b>A{idx}:</b> {a}", body_style))
                story.append(Spacer(1, 6))

        # Quiz
        if note.mcqs:
            story.append(Spacer(1, 10))
            story.append(Paragraph("Multiple Choice Quiz", h1_style))
            for idx, mcq in enumerate(note.mcqs, 1):
                q = _inline_md(str(mcq.get('question', '')))
                story.append(Paragraph(f"<b>Question {idx}:</b> {q}", body_style))
                for opt in mcq.get('options', []):
                    story.append(Paragraph(f"&nbsp;&nbsp;○&nbsp;{_inline_md(str(opt))}", bullet_style))
                ans = _inline_md(str(mcq.get('answer', '')))
                exp = _inline_md(str(mcq.get('explanation', '')))
                story.append(Paragraph(f"<i>✓ Correct: {ans}</i>", body_style))
                story.append(Paragraph(f"<i>Explanation: {exp}</i>", body_style))
                story.append(Spacer(1, 8))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

export_service = ExportService()
